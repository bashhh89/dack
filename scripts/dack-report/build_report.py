#!/usr/bin/env python3
import argparse
import html
import json
import re
import shutil
import subprocess
import tempfile
from collections import OrderedDict
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "12324A"
BLUE = "2F5D7C"
GOLD = "B68A35"
LIGHT_BLUE = "EAF1F6"
LIGHT_GOLD = "F7F0E4"
LIGHT_GRAY = "F4F6F8"
TEXT = "20252B"
WHITE = "FFFFFF"


METRIC_PATTERNS = [
    ("Invitations sent", r"([\d,]+)\s+Invitations?\s+sent"),
    ("Contacts reached", r"([\d,]+)\s+Total\s+Contacts?\s+reached"),
    ("Personal outreach emails", r"([\d,]+)\s+Personal\s+Outreach\s+Emails?"),
    ("Page views", r"([\d,]+)\s+Page\s+views?"),
    ("Opened", r"([\d,]+)\s+Opened"),
    ("Eventbrite registrants", r"([\d,]+)\s+Eventbrite\s+Registrants?"),
    ("Total attendees", r"([\d,]+)\s+Total\s+Attendees?"),
    ("Additional walk-in attendees", r"([\d,]+)\s+Additional\s+Walk[- ]In\s+Attendees?"),
    ("Business cards collected", r"([\d,]+)\s+Business\s+Cards?\s+Collected"),
    ("Notice of Interest submissions", r"([\d,]+)\s+Notice\s+of\s+Interest\s+Submissions?|([\d,]+)\s+NOI"),
    ("MWBE/SDVOB participation goal", r"Total\s+MWBE/SDVOB\s+participation\s+goal:\s*\$([\d,]+)"),
    ("Awarded to MWBE/SDVOB firms", r"Awarded\s+to\s+MWBE/SDVOB\s+firms\s+to\s+date:\s*\$([\d,]+)"),
    ("Toward-goal participation", r"Towards\s+the\s+goal\s+with\s+60%\s+for\s+Suppliers:\s*\$([\d,]+)"),
    ("Small business firms", r"Small\s+Business\s+[-–]\s+([\d,]+)\s+firms"),
    ("MWBE/SDVOB payments to date", r"Total\s+MWBE/SDVOB\s+payments\s+to\s+date:\s*\$([\d,]+)"),
    ("Q1 MWBE/SDVOB payments", r"Q1-specific\s+MWBE/SDVOB\s+payments:\s*\$([\d,]+)"),
    ("Certified payrolls received", r"During\s+the\s+reporting\s+period,\s+([\d,]+)\s+certified\s+payrolls\s+were\s+received"),
    ("Labor-related site visits", r"([\d,]+)\s+labor-related\s+site\s+visits\s+were\s+conducted"),
    ("Workforce diversity site visits", r"conducted\s+three\s+\(([\d,]+)\)\s+workforce\s+diversity\s+site\s+visits"),
]


def clean_text(value):
    value = str(value or "")
    value = value.replace("\u2013", "-").replace("\u2014", "-")
    value = value.replace("\u201c", '"').replace("\u201d", '"')
    value = value.replace("\u2018", "'").replace("\u2019", "'")
    return re.sub(r"\s+", " ", value).strip()


def read_source(path):
    source = Path(path)
    if not source.exists():
        raise FileNotFoundError(source)
    if source.suffix.lower() == ".pdf":
        pdftotext = shutil.which("pdftotext")
        if not pdftotext:
            raise RuntimeError("PDF source requires pdftotext, but it is not installed.")
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        try:
            subprocess.run([pdftotext, "-layout", str(source), str(tmp_path)], check=True)
            text = tmp_path.read_text(errors="ignore")
        finally:
            tmp_path.unlink(missing_ok=True)
        paragraphs = [clean_text(x) for x in re.split(r"\n\s*\n|\n", text) if clean_text(x)]
        return paragraphs, []
    if source.suffix.lower() == ".docx":
        doc = Document(source)
        paragraphs = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                tables.append(rows)
        return paragraphs, tables
    text = source.read_text(errors="ignore")
    paragraphs = [clean_text(x) for x in re.split(r"\n\s*\n|\n", text) if clean_text(x)]
    return paragraphs, []


def extract_between(paragraphs, heading, stop_headings):
    start = None
    for i, paragraph in enumerate(paragraphs):
        if paragraph.strip().lower() == heading.lower():
            start = i + 1
            break
    if start is None:
        return []
    out = []
    for paragraph in paragraphs[start:]:
        normalized = paragraph.strip().lower()
        if any(normalized == h.lower() for h in stop_headings):
            break
        out.append(paragraph)
    return out


def extract_metrics(paragraphs):
    blob = "\n".join(paragraphs)
    metrics = OrderedDict()
    for label, pattern in METRIC_PATTERNS:
        match = re.search(pattern, blob, re.IGNORECASE)
        if not match:
            continue
        value = next((g for g in match.groups() if g), None)
        if value:
            metrics[label] = int(value.replace(",", ""))
    opened_rate = re.search(r"Opened\s*\(([\d.]+%)\)", blob, re.IGNORECASE)
    if opened_rate:
        metrics["Open rate"] = opened_rate.group(1)
    return metrics


def extract_project_context(paragraphs):
    context = []
    for paragraph in paragraphs:
        lower = paragraph.lower()
        if any(term in lower for term in ["commits to hiring", "was engaged", "retained dack", "workforce diversity site visits"]):
            context.append(paragraph)
        if len(context) >= 4:
            break
    return context


def extract_events(paragraphs):
    events = []
    for paragraph in paragraphs:
        lower = paragraph.lower()
        if "report located in the exhibits" in lower or "event outreach" in lower:
            events.append(paragraph)
    return events[:6]


def extract_goals(paragraphs):
    goals = extract_between(paragraphs, "GOALS", ["Outreach", "COMPLIANCE", "Compliance"])
    return [g for g in goals if len(g) > 25][:8]


def find_first(patterns, paragraphs, default="Needs confirmation"):
    blob = "\n".join(paragraphs)
    for pattern in patterns:
        match = re.search(pattern, blob, re.IGNORECASE)
        if match:
            return clean_text(match.group(1) if match.groups() else match.group(0))
    return default


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(clean_text(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_horizontal_rule(paragraph, color=GOLD):
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    p_pr.append(border)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)

    for style_name, size, color in [
        ("Title", 24, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 12, BLUE),
        ("Heading 3", 10, NAVY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_section_heading(doc, title, eyebrow=None):
    if eyebrow:
        p = doc.add_paragraph()
        r = p.add_run(eyebrow.upper())
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(GOLD)
    h = doc.add_heading(title, level=1)
    add_horizontal_rule(h)


def add_cover(doc, title, client, period, prepared_by):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    cell.height = Inches(2.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(WHITE)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(client)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("DDE8EF")

    doc.add_paragraph("")
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    labels = [
        ("Reporting Period", period),
        ("Prepared By", prepared_by),
        ("Document Type", "Executive Compliance Report"),
        ("Draft Status", "Client-ready draft based on provided source material"),
    ]
    for row, (label, value) in zip(meta.rows, labels):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_text(row.cells[0], label, bold=True, color=NAVY, size=9)
        set_cell_text(row.cells[1], value, size=9)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared for executive, compliance, governance, and stakeholder review.")
    r.font.italic = True
    r.font.size = Pt(10)


def add_highlight_cards(doc, highlights):
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, item in enumerate(highlights[:6]):
        cell = table.rows[idx // 3].cells[idx % 3]
        set_cell_shading(cell, LIGHT_GOLD if idx % 2 else LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        metric = p.add_run(str(item["value"]))
        metric.bold = True
        metric.font.size = Pt(17)
        metric.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run("\n")
        label = p.add_run(item["label"])
        label.font.size = Pt(8.5)
        label.font.color.rgb = RGBColor.from_string(TEXT)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, header, bold=True, color=WHITE, size=8.5)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            set_cell_text(row.cells[i], value, size=8.3)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def add_simple_toc(doc, sections):
    add_section_heading(doc, "Table of Contents", "Report navigation")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_shading(table.rows[0].cells[0], NAVY)
    set_cell_shading(table.rows[0].cells[1], NAVY)
    set_cell_text(table.rows[0].cells[0], "Section", bold=True, color=WHITE, size=8.5)
    set_cell_text(table.rows[0].cells[1], "Purpose", bold=True, color=WHITE, size=8.5)
    for section, purpose in sections:
        row = table.add_row()
        set_cell_text(row.cells[0], section, bold=True, color=NAVY, size=8.5)
        set_cell_text(row.cells[1], purpose, size=8.2)


def make_chart(metrics, output_dir):
    numeric_items = [(k, v) for k, v in metrics.items() if isinstance(v, int)]
    if len(numeric_items) < 2:
        return None
    labels = [k for k, _ in numeric_items[:8]]
    values = [v for _, v in numeric_items[:8]]
    fig, ax = plt.subplots(figsize=(7.4, 3.8))
    bars = ax.barh(labels[::-1], values[::-1], color="#2F5D7C")
    ax.set_title("Confirmed Outreach Activity Metrics", fontsize=13, weight="bold", color="#12324A")
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.tick_params(axis="y", labelsize=8)
    ax.tick_params(axis="x", labelsize=8)
    ax.grid(axis="x", alpha=0.18)
    for bar in bars:
        width = bar.get_width()
        ax.text(width, bar.get_y() + bar.get_height() / 2, f" {int(width):,}", va="center", fontsize=8)
    fig.tight_layout()
    chart_path = Path(output_dir) / "outreach-metrics.png"
    fig.savefig(chart_path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return chart_path


def html_rows(headers, rows):
    head = "".join(f"<th>{html.escape(str(h))}</th>" for h in headers)
    body = []
    for row in rows:
        body.append("<tr>" + "".join(f"<td>{html.escape(str(cell))}</td>" for cell in row) + "</tr>")
    return f"<table><thead><tr>{head}</tr></thead><tbody>{''.join(body)}</tbody></table>"


def render_pdf_with_chrome(html_path, pdf_path):
    chrome = shutil.which("google-chrome") or shutil.which("chromium") or shutil.which("chromium-browser")
    if not chrome:
        raise RuntimeError("PDF output requires google-chrome/chromium for headless rendering.")
    subprocess.run(
        [
            chrome,
            "--headless=new",
            "--disable-gpu",
            "--no-sandbox",
            "--no-pdf-header-footer",
            f"--print-to-pdf={pdf_path}",
            str(html_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def write_pdf_report(args, paragraphs, tables, metrics, goals, events, context, output_docx, pdf_output):
    client = args.client or find_first([r"(Resorts World New York City)", r"(RWNYC)", r"(Brookfield Commons[^.]+)", r"(NYPA HEADQUARTERS PROJECT)"], paragraphs)
    period = args.period or find_first([r"(January 1, 2026,?\s+to\s+March 31, 2026)", r"(Q1 2026)", r"(2025 reporting period)"], paragraphs)
    title = args.title or f"{client} Compliance Report"
    prepared_by = args.prepared_by

    metric_cards = []
    for label, value in metrics.items():
        metric_cards.append(
            f"<div class='metric'><strong>{html.escape(f'{value:,}' if isinstance(value, int) else str(value))}</strong><span>{html.escape(label)}</span></div>"
        )
    if not metric_cards:
        metric_cards.append("<div class='metric'><strong>Needs confirmation</strong><span>Confirmed source metrics</span></div>")

    goal_rows = [
        ["MWBE goal", find_first([r"(\d+% MWBE/SDVOB participation goal)", r"(\d+%\s+MWBE\s+participation)", r"(\d+%\s+MWBE\s+utilization\s+goal)"], paragraphs), "Confirm against project agreement or utilization plan."],
        ["MBE goal", find_first([r"(\d+(?:\.\d+)?%\s+MBE)", r"(15%\s+Minority Business Enterprise)"], paragraphs), "Needs source-backed dollar target if not provided."],
        ["WBE goal", find_first([r"(\d+(?:\.\d+)?%\s+WBE)", r"(15%\s+Women Business Enterprise)"], paragraphs), "Needs source-backed dollar target if not provided."],
        ["SDVOB goal", find_first([r"(\d+(?:\.\d+)?%\s+SDVOB)", r"(6%\s+Service-Disabled Veteran-Owned Business)"], paragraphs), "Needs source-backed dollar target if not provided."],
    ]

    utilization_rows = [
        ["Awards / utilization plan", find_first([r"(Awarded to MWBE/SDVOB firms to date:\s*\$[\d,]+)", r"(total awards counted towards the goal.*?\$[\d,]+)"], paragraphs), "Use awarded contracts, utilization plans, POs, and certification records."],
        ["Supplier adjustment", find_first([r"(60%[^.]+suppliers)", r"(60% correction for suppliers)"], paragraphs), "Apply only when source says supplier credit is adjusted."],
        ["MBE awards / toward goal", find_first([r"(MBE Awards:\s*\$[\d,]+)", r"(Minority-Owned Business Enterprises.*?\$[\d,]+)"], paragraphs), "Needs confirmation if not provided."],
        ["WBE awards / toward goal", find_first([r"(WBE Awards:\s*\$[\d,]+)", r"(Women-Owned Business Enterprises.*?\$[\d,]+)"], paragraphs), "Needs confirmation if not provided."],
        ["SDVOB awards / toward goal", find_first([r"(SDVOB Awards:\s*\$[\d,]+)", r"(Service-Disabled Veteran-Owned Businesses.*?\$[\d,]+)"], paragraphs), "Needs confirmation if not provided."],
    ]

    payment_rows = [
        ["Paid to date", find_first([r"(Total MWBE/SDVOB payments to date:\s*\$[\d,]+)", r"(\$[\d,]+ was paid to MWBE/SDVOB)"], paragraphs), "Confirm against payment logs and check registers."],
        ["Quarter-specific payments", find_first([r"(Q1-specific MWBE/SDVOB payments:\s*\$[\d,]+)", r"(Payments Reported\s+In This Period.*?\$[\d,]+)"], paragraphs), "Use only when source clearly provides period payment."],
        ["Percent paid", find_first([r"(Percentage of total MWBE/SDVOB contract value paid:\s*\d+%)", r"(representing about \d+% of the total contract amounts)"], paragraphs), "Do not calculate if denominator is missing."],
    ]

    workforce_rows = [
        ["Total construction hours", find_first([r"(Total construction hours:\s*[\d,.]+ hours)", r"(Total Work Hours:\s*[\d,.]+ construction hours)"], paragraphs), "Certified payroll / LCP Tracker source required."],
        ["Minority workforce participation", find_first([r"(Minority workforce participation:\s*[\d.]+%)", r"(Minority Workforce.*?[\d.]+% of total work hours)"], paragraphs), "May be revised as payroll data is finalized."],
        ["Female workforce participation", find_first([r"(Female workforce participation:\s*[\d.]+%)", r"(Female Workforce.*?[\d.]+% of total work hours)"], paragraphs), "May be revised as payroll data is finalized."],
    ]

    labor_rows = [
        ["Certified payrolls received", find_first([r"(During the reporting period,\s*\d+\s+certified payrolls were received)", r"(\d+\s+certified payrolls were received)"], paragraphs), "Use exact count from source."],
        ["Labor-related site visits", find_first([r"(\d+\s+labor-related site visits were conducted)", r"(three\s+\(\d+\)\s+workforce diversity site visits)"], paragraphs), "Use exact count from source."],
        ["System review", find_first([r"(LCP Tracker[^.]+discrepancies[^.]+)", r"(LCP Tracker[^.]+approved[^.]+)"], paragraphs), "Do not state no discrepancies unless source supports it."],
    ]

    event_list = "".join(f"<li>{html.escape(e)}</li>" for e in events[:6]) or "<li>Needs confirmation: outreach events were not clearly extractable from the source.</li>"
    context_paras = "".join(f"<p>{html.escape(p)}</p>" for p in context[:4]) or "<p>Needs confirmation: project context was not clearly identified in the source material.</p>"
    appendix_rows = [[f"Table {idx}", str(len(table)), " | ".join(table[0][:4]) if table else "Needs confirmation"] for idx, table in enumerate(tables[:8], start=1)]

    html_text = f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<title>{html.escape(title)}</title>
<style>
@page {{ size: Letter; margin: 0.65in 0.7in; }}
body {{ font-family: Arial, Helvetica, sans-serif; color: #20252B; font-size: 10.5pt; line-height: 1.55; }}
.cover {{ page-break-after: always; padding-top: 1.2in; }}
.band {{ background: #12324A; color: white; padding: 28px 34px; margin: 0 -0.2in 36px -0.2in; }}
.band h1 {{ margin: 0; font-size: 28pt; letter-spacing: 0.3px; text-transform: uppercase; }}
.meta {{ margin-top: 28px; width: 70%; }}
h2 {{ font-size: 20pt; text-transform: uppercase; margin: 30px 0 10px; border-bottom: 2px solid #B68A35; padding-bottom: 6px; }}
h3 {{ font-size: 12pt; margin: 18px 0 6px; }}
.eyebrow {{ color: #B68A35; font-size: 8pt; font-weight: bold; text-transform: uppercase; margin-top: 18px; }}
.metrics {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 10px; margin: 14px 0 18px; }}
.metric {{ background: #EAF1F6; padding: 13px; min-height: 58px; }}
.metric strong {{ display: block; color: #12324A; font-size: 18pt; }}
.metric span {{ display: block; font-size: 8.5pt; }}
table {{ border-collapse: collapse; width: 100%; margin: 12px 0 18px; page-break-inside: avoid; }}
th {{ background: #12324A; color: white; text-align: left; font-size: 8.5pt; padding: 7px; }}
td {{ border: 1px solid #CBD5DD; padding: 7px; vertical-align: top; font-size: 8.5pt; }}
tr:nth-child(even) td {{ background: #F4F6F8; }}
ul {{ margin-top: 6px; }}
.footer-note {{ color: #5A6570; font-size: 8.5pt; margin-top: 24px; }}
</style>
</head>
<body>
<section class="cover">
  <div class="band"><h1>{html.escape(title)}</h1></div>
  <h3>{html.escape(client)}</h3>
  <table class="meta">
    <tr><th>Reporting Period</th><td>{html.escape(period)}</td></tr>
    <tr><th>Prepared By</th><td>{html.escape(prepared_by)}</td></tr>
    <tr><th>Document Type</th><td>MWBE/EEO/Labor Compliance Report</td></tr>
  </table>
</section>
<h2>Table of Contents</h2>
{html_rows(["Section", "Purpose"], [["Executive Summary", "Quarterly compliance status and key metrics."], ["Project Goals", "Participation framework and eligible spend."], ["MWBE / SDVOB Outreach", "Good-faith outreach record."], ["MWBE / SDVOB Utilization", "Awards and goal progress."], ["MWBE / SDVOB Payments Monitoring", "Payment accountability."], ["Workforce Diversity", "Payroll, MEUR, and workforce participation."], ["Labor Compliance", "Prevailing wage and site oversight."], ["Attachments", "Audit-supporting source records."]])}
<h2>Executive Summary</h2>
<p>This Quarterly MWBE/EEO/Labor report summarizes outreach, utilization, payments, workforce diversity, and labor compliance activity for {html.escape(client)} during {html.escape(period)}. {html.escape(prepared_by)} supports the project team through compliance monitoring, documentation review, outreach coordination, and reporting support.</p>
<p>The source material reflects a compliance program focused on good-faith outreach, certified-firm participation, workforce visibility, payment accountability, and audit-ready documentation. Confirmed metrics are presented only where the source provides support.</p>
<div class="metrics">{''.join(metric_cards[:9])}</div>
<h2>Project Goals</h2>
{html_rows(["Goal Area", "Source Value", "Reporting Note"], goal_rows)}
<h2>Project / Program Context</h2>
{context_paras}
<h2>MWBE / SDVOB Outreach</h2>
<p>Outreach should document how eligible MWBE and SDVOB firms were identified, engaged, and tracked.</p>
<ul>{event_list}</ul>
<h2>MWBE / SDVOB Utilization</h2>
{html_rows(["Analysis Area", "Current Source Value", "Reporting Note"], utilization_rows)}
<h2>MWBE / SDVOB Payments Monitoring</h2>
{html_rows(["Payment Area", "Source Value", "Reporting Note"], payment_rows)}
<h2>Workforce Diversity</h2>
{html_rows(["Workforce Area", "Source Value", "Reporting Note"], workforce_rows)}
<h2>Labor Compliance</h2>
{html_rows(["Labor Area", "Source Value", "Reporting Note"], labor_rows)}
<h2>Needs Confirmation</h2>
<ol>
  <li>Confirm all placeholder values, dollar goals, and participation percentages.</li>
  <li>Attach or cross-reference exhibits, flyers, event reports, outreach logs, payment records, certified payroll totals, MEURs, and site-visit reports.</li>
  <li>Add utilization and workforce charts only after final data is verified.</li>
</ol>
<h2>Attachments</h2>
{html_rows(["Source Table", "Rows", "First Row / Description"], appendix_rows) if appendix_rows else "<p>Needs confirmation: no source tables were extracted.</p>"}
<p class="footer-note">Generated alongside Word file: {html.escape(str(output_docx))}</p>
</body>
</html>"""

    pdf_output = Path(pdf_output).expanduser().resolve()
    pdf_output.parent.mkdir(parents=True, exist_ok=True)
    html_output = Path(args.html_output).expanduser().resolve() if args.html_output else pdf_output.with_suffix(".html")
    html_output.parent.mkdir(parents=True, exist_ok=True)
    html_output.write_text(html_text, encoding="utf-8")
    render_pdf_with_chrome(html_output, pdf_output)
    return html_output, pdf_output


def build_report(args):
    paragraphs, tables = read_source(args.source)
    metrics = extract_metrics(paragraphs)
    goals = extract_goals(paragraphs)
    events = extract_events(paragraphs)
    context = extract_project_context(paragraphs)

    client = args.client or find_first([r"(Resorts World New York City)", r"(RWNYC)", r"(Brookfield Commons[^.]+)"], paragraphs)
    period = args.period or find_first([r"(January 1, 2026,?\s+to\s+March 31, 2026)", r"(Q1 2026)", r"(2025 reporting period)"], paragraphs)
    title = args.title or f"{client} Compliance Report"
    prepared_by = args.prepared_by

    doc = Document()
    configure_document(doc)
    add_cover(doc, title, client, period, prepared_by)
    doc.add_section(WD_SECTION.NEW_PAGE)

    toc_sections = [
        ("Executive Summary", "Quarterly compliance status, key metrics, and DACK role."),
        ("Project Goals", "MWBE/SDVOB targets, eligible spend logic, and good-faith framework."),
        ("MWBE / SDVOB Outreach", "Documented outreach channels, events, and bid-log evidence."),
        ("MWBE / SDVOB Utilization", "Awards, utilization plans, certification checks, and goal progress."),
        ("MWBE / SDVOB Payments Monitoring", "Invoice/payment review and paid-to-date evidence."),
        ("Workforce Diversity", "Certified payroll, MEUR, workforce hours, and contractor-level participation."),
        ("Labor Compliance", "Prevailing wage, certified payroll, classification, and site-visit oversight."),
        ("Attachments", "Audit-supporting logs, reports, payment records, and source tables."),
    ]
    add_simple_toc(doc, toc_sections)
    doc.add_section(WD_SECTION.NEW_PAGE)

    add_section_heading(doc, "Executive Summary", "Overview")
    summary_paragraphs = [
        f"This Quarterly MWBE/EEO/Labor report summarizes outreach, utilization, payments, workforce diversity, and labor compliance activity for {client} during {period}. {prepared_by} supports the project team through compliance monitoring, documentation review, outreach coordination, and reporting support.",
        "The source material reflects a compliance program focused on good-faith outreach, certified-firm participation, workforce visibility, payment accountability, and audit-ready documentation. Confirmed metrics are presented only where the source provides support.",
        "Outreach activity is documented through source-backed engagement metrics, event references, bid-log style records, and follow-up indicators. These records help demonstrate that eligible firms were identified, contacted, and tracked as part of the project's participation strategy.",
        "Items requiring additional data are labeled as Needs confirmation so the report can remain useful as a client-ready draft without overstating unsupported participation, payment, or workforce claims.",
    ]
    for text in summary_paragraphs:
        doc.add_paragraph(text)

    add_section_heading(doc, "Current Participation and Activity Metrics", "Confirmed source-backed items")
    highlight_items = []
    for label, value in list(metrics.items())[:6]:
        highlight_items.append({"label": label, "value": f"{value:,}" if isinstance(value, int) else value})
    while len(highlight_items) < 6:
        highlight_items.append({"label": "Additional metric", "value": "Needs confirmation"})
    add_highlight_cards(doc, highlight_items)

    add_section_heading(doc, "Project Goals", "Participation framework")
    goal_rows = [
        ["MWBE goal", find_first([r"(\d+% MWBE/SDVOB participation goal)", r"(\d+%\s+MWBE\s+participation)", r"(\d+%\s+MWBE\s+utilization\s+goal)"], paragraphs), "Confirm against project agreement or utilization plan."],
        ["MBE goal", find_first([r"(\d+(?:\.\d+)?%\s+MBE)", r"(15%\s+Minority Business Enterprise)"], paragraphs), "Needs source-backed dollar target if not provided."],
        ["WBE goal", find_first([r"(\d+(?:\.\d+)?%\s+WBE)", r"(15%\s+Women Business Enterprise)"], paragraphs), "Needs source-backed dollar target if not provided."],
        ["SDVOB goal", find_first([r"(\d+(?:\.\d+)?%\s+SDVOB)", r"(6%\s+Service-Disabled Veteran-Owned Business)"], paragraphs), "Needs source-backed dollar target if not provided."],
        ["Eligible / subjected value", find_first([r"(approximately \$[\d,.]+ million is subject to MWBE/SDVOB participation)", r"(\$[\d,]+(?:\.\d+)?\s+.*subject(?:ed)? to participation)"], paragraphs), "Use only if source identifies the eligible denominator."],
    ]
    add_table(doc, ["Goal Area", "Source Value", "Reporting Note"], goal_rows)

    add_section_heading(doc, "Project / Program Context", "Background")
    if context:
        for text in context[:4]:
            doc.add_paragraph(text)
    else:
        doc.add_paragraph("Needs confirmation: project context was not clearly identified in the source material.")

    add_section_heading(doc, "DACK Role and Compliance Activities", "Commitments and status")
    scope_rows = []
    if goals:
        for goal in goals[:6]:
            scope_rows.append([
                goal,
                "Activity documented in source materials",
                "Supports good faith effort documentation, reporting transparency, and audit readiness.",
                "In progress / continue documentation",
            ])
    else:
        scope_rows.append([
            "Needs confirmation",
            "Source did not provide a distinct scope list.",
            "Needs confirmation",
            "Confirm DACK contracted scope and deliverables.",
        ])
    add_table(
        doc,
        ["Scope Commitment", "Activity Completed", "Compliance / Business Outcome", "Status or Follow-Up"],
        scope_rows,
        [1.8, 1.7, 2.1, 1.4],
    )

    add_section_heading(doc, "MWBE / SDVOB Outreach", "Engagement record")
    doc.add_paragraph(
        "Outreach should document how eligible MWBE and SDVOB firms were identified, engaged, and tracked. The outreach record may include informational blasts, social media postings, chambers and civic organizations, contractor lists organized by trade or scope, informational sessions, bid logs, and upcoming supplier diversity events."
    )
    if events:
        for event in events:
            doc.add_paragraph(event, style="List Bullet")
    else:
        doc.add_paragraph("Needs confirmation: no distinct outreach events were extracted from the source.")

    metric_rows = []
    for label, value in metrics.items():
        metric_rows.append([label, f"{value:,}" if isinstance(value, int) else str(value), "Confirmed in source text"])
    if metric_rows:
        add_table(doc, ["Outreach Metric", "Value", "Source Status"], metric_rows)
    else:
        doc.add_paragraph("Needs confirmation: source did not contain chart-ready outreach metrics.")

    with tempfile.TemporaryDirectory() as td:
        chart = make_chart(metrics, td)
        if chart:
            doc.add_paragraph("")
            doc.add_picture(str(chart), width=Inches(6.8))

    add_section_heading(doc, "MWBE / SDVOB Utilization", "Awards and goal progress")
    analysis_rows = [
        ["Awards / utilization plan", find_first([r"(Awarded to MWBE/SDVOB firms to date:\s*\$[\d,]+)", r"(total awards counted towards the goal.*?\$[\d,]+)"], paragraphs), "Use awarded contracts, utilization plans, POs, and certification records."],
        ["Supplier adjustment", find_first([r"(60%[^.]+suppliers)", r"(60% correction for suppliers)"], paragraphs), "Apply only when source says supplier credit is adjusted."],
        ["MBE awards / toward goal", find_first([r"(MBE Awards:\s*\$[\d,]+)", r"(Minority-Owned Business Enterprises.*?\$[\d,]+)"], paragraphs), "Needs confirmation if not provided."],
        ["WBE awards / toward goal", find_first([r"(WBE Awards:\s*\$[\d,]+)", r"(Women-Owned Business Enterprises.*?\$[\d,]+)"], paragraphs), "Needs confirmation if not provided."],
        ["SDVOB awards / toward goal", find_first([r"(SDVOB Awards:\s*\$[\d,]+)", r"(Service-Disabled Veteran-Owned Businesses.*?\$[\d,]+)"], paragraphs), "Needs confirmation if not provided."],
    ]
    add_table(doc, ["Analysis Area", "Current Source Value", "Reporting Note"], analysis_rows)

    add_section_heading(doc, "MWBE / SDVOB Payments Monitoring", "Payment accountability")
    payment_rows = [
        ["Paid to date", find_first([r"(Total MWBE/SDVOB payments to date:\s*\$[\d,]+)", r"(\$[\d,]+ was paid to MWBE/SDVOB)"], paragraphs), "Confirm against payment logs and check registers."],
        ["Quarter-specific payments", find_first([r"(Q1-specific MWBE/SDVOB payments:\s*\$[\d,]+)", r"(Payments Reported\s+In This Period.*?\$[\d,]+)"], paragraphs), "Use only when source clearly provides period payment."],
        ["Percent paid", find_first([r"(Percentage of total MWBE/SDVOB contract value paid:\s*\d+%)", r"(representing about \d+% of the total contract amounts)"], paragraphs), "Do not calculate if denominator is missing."],
        ["Verification method", "Invoice review, contractor payment records, and approved utilization plans where provided.", "Preserve supporting records as attachments."],
    ]
    add_table(doc, ["Payment Area", "Source Value", "Reporting Note"], payment_rows)

    add_section_heading(doc, "Workforce Diversity", "Payroll and MEUR analysis")
    workforce_rows = [
        ["Total construction hours", find_first([r"(Total construction hours:\s*[\d,.]+ hours)", r"(Total Work Hours:\s*[\d,.]+ construction hours)"], paragraphs), "Certified payroll / LCP Tracker source required."],
        ["Minority workforce participation", find_first([r"(Minority workforce participation:\s*[\d.]+%)", r"(Minority Workforce.*?[\d.]+% of total work hours)"], paragraphs), "May be revised as payroll data is finalized."],
        ["Female workforce participation", find_first([r"(Female workforce participation:\s*[\d.]+%)", r"(Female Workforce.*?[\d.]+% of total work hours)"], paragraphs), "May be revised as payroll data is finalized."],
        ["Contractor-level table", "Needs confirmation" if "CONTRACTOR NAME" not in "\n".join(paragraphs) else "Contractor-level minority/female participation table present in source.", "Include when available."],
    ]
    add_table(doc, ["Workforce Area", "Source Value", "Reporting Note"], workforce_rows)

    add_section_heading(doc, "Labor Compliance", "Prevailing wage and site oversight")
    labor_rows = [
        ["Certified payrolls received", find_first([r"(During the reporting period,\s*\d+\s+certified payrolls were received)", r"(\d+\s+certified payrolls were received)"], paragraphs), "Use exact count from source."],
        ["Labor-related site visits", find_first([r"(\d+\s+labor-related site visits were conducted)", r"(three\s+\(\d+\)\s+workforce diversity site visits)"], paragraphs), "Use exact count from source."],
        ["System review", find_first([r"(LCP Tracker[^.]+discrepancies[^.]+)", r"(LCP Tracker[^.]+approved[^.]+)"], paragraphs), "Do not state no discrepancies unless source supports it."],
        ["Monitoring activities", "Certified payroll review, wage/classification verification, manpower reconciliation, and random site visits.", "Include only activities supported by source or mark Needs confirmation."],
    ]
    add_table(doc, ["Labor Area", "Source Value", "Reporting Note"], labor_rows)

    add_section_heading(doc, "Conclusion", "Executive close")
    doc.add_paragraph(
        f"The available source material supports a client-ready draft report showing DACK's role in outreach, utilization monitoring, payment review, workforce diversity tracking, and labor compliance for {client}. The report should now be reviewed against the underlying records, with missing participation, workforce, and payment data confirmed before final distribution."
    )

    add_section_heading(doc, "Needs Confirmation", "Data gaps")
    confirmation_items = [
        "Confirm all placeholder values, dollar goals, and participation percentages.",
        "Attach or cross-reference exhibits, flyers, event reports, and outreach logs.",
        "Add utilization and workforce charts only after final data is verified.",
        "Confirm payment logs, certified payroll totals, MEURs, and site-visit reports before final issue.",
    ]
    for step in confirmation_items:
        doc.add_paragraph(step, style="List Number")

    if tables:
        add_section_heading(doc, "Attachments", "Audit backup")
        appendix_rows = []
        for idx, table in enumerate(tables[:8], start=1):
            first = " | ".join(table[0][:4]) if table else ""
            appendix_rows.append([f"Table {idx}", str(len(table)), first or "Needs confirmation"])
        add_table(doc, ["Source Table", "Rows", "First Row / Description"], appendix_rows)

    output = Path(args.output).expanduser().resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    html_output = None
    pdf_output = None
    if args.pdf_output:
        html_output, pdf_output = write_pdf_report(
            args, paragraphs, tables, metrics, goals, events, context, output, args.pdf_output
        )

    manifest = {
        "output": str(output),
        "pdf_output": str(pdf_output) if pdf_output else None,
        "html_output": str(html_output) if html_output else None,
        "paragraphs_read": len(paragraphs),
        "tables_read": len(tables),
        "metrics": metrics,
        "needs_confirmation": [
            "Final contract dollar goals",
            "Final MWBE/SDVOB/LBE utilization percentages",
            "Worker-hour and demographic totals",
            "Payment verification totals",
            "Final exhibit list and sign-off owner",
        ],
    }
    print(json.dumps(manifest, indent=2))


def main():
    parser = argparse.ArgumentParser(description="Build a polished DACK executive compliance DOCX report.")
    parser.add_argument("--source", required=True, help="Source rough report, notes, or text file.")
    parser.add_argument("--reference", help="Optional reference report for style benchmark. Not used as factual source.")
    parser.add_argument("--output", required=True, help="Output .docx path.")
    parser.add_argument("--pdf-output", help="Optional output .pdf path rendered with local headless Chrome.")
    parser.add_argument("--html-output", help="Optional intermediate .html path used for PDF rendering.")
    parser.add_argument("--title", help="Report title.")
    parser.add_argument("--client", help="Client/project name.")
    parser.add_argument("--period", help="Reporting period.")
    parser.add_argument("--prepared-by", default="DACK Consulting Solutions, Inc.")
    args = parser.parse_args()
    build_report(args)


if __name__ == "__main__":
    main()
